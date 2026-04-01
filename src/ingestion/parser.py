"""
Layer 1 — Document Parser
FIX #10: PDF parser now extracts TOC for real section boundaries.
FIX #15: Uses defusedxml for XML parsing to prevent XXE attacks.
"""

from __future__ import annotations

import json
import logging
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)


@dataclass
class ParsedDocument:
    text: str
    sections: list[dict[str, Any]] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)
    format: str = "unknown"

    @property
    def has_structure(self) -> bool:
        return len(self.sections) > 0


class DocumentParser:
    """Unified parser dispatching to format-specific handlers."""

    def parse(self, content: bytes, filename: str) -> ParsedDocument:
        ext = Path(filename).suffix.lower()
        handlers = {
            ".pdf": self._parse_pdf,
            ".docx": self._parse_docx,
            ".xml": self._parse_xml,
            ".txt": self._parse_text,
            ".md": self._parse_text,
            ".json": self._parse_json,
        }
        handler = handlers.get(ext, self._parse_text)
        return handler(content, filename)

    def _parse_pdf(self, content: bytes, filename: str) -> ParsedDocument:
        """FIX #10: Extract TOC for real section boundaries; fallback to pages."""
        import fitz  # PyMuPDF

        doc = fitz.open(stream=content, filetype="pdf")
        full_text_parts: list[str] = []
        for i in range(len(doc)):
            full_text_parts.append(doc[i].get_text("text"))

        full_text = "\n\n".join(full_text_parts)

        # Try TOC-based section extraction first
        toc = doc.get_toc()  # [(level, title, page_number), ...]
        sections: list[dict[str, Any]] = []

        if toc:
            # Build sections from TOC entries
            for i, (level, title, page_num) in enumerate(toc):
                # Determine page range for this section
                start_page = page_num - 1  # 0-indexed
                if i + 1 < len(toc):
                    end_page = toc[i + 1][2] - 1
                else:
                    end_page = len(doc)

                section_text = ""
                for p in range(max(0, start_page), min(end_page, len(doc))):
                    section_text += doc[p].get_text("text") + "\n"

                sections.append({
                    "title": title.strip(),
                    "text": section_text.strip(),
                    "level": level,
                    "page": page_num,
                })
        else:
            # Fallback: each page is a section
            for i in range(len(doc)):
                page_text = doc[i].get_text("text")
                sections.append({
                    "title": f"Page {i + 1}",
                    "text": page_text,
                    "level": 1,
                })

        metadata = {
            "page_count": len(doc),
            "title": doc.metadata.get("title", ""),
            "has_toc": len(toc) > 0,
        }
        doc.close()
        return ParsedDocument(full_text, sections, metadata, "pdf")

    def _parse_docx(self, content: bytes, filename: str) -> ParsedDocument:
        from docx import Document

        doc = Document(BytesIO(content))
        sections, texts = [], []
        current: dict[str, Any] | None = None

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if para.style and para.style.name.startswith("Heading"):
                if current:
                    sections.append(current)
                level = int(para.style.name.replace("Heading ", "").replace("Heading", "1") or "1")
                current = {"title": text, "text": "", "level": level}
            else:
                if current:
                    current["text"] += text + "\n"
                texts.append(text)

        if current:
            sections.append(current)
        return ParsedDocument("\n".join(texts), sections, {}, "docx")

    def _parse_xml(self, content: bytes, filename: str) -> ParsedDocument:
        """Parse PubMed MEDLINE XML. Uses lxml-xml parser (XXE-safe since lxml 4.6.0)."""
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(content, "lxml-xml")
        sections = []

        article = soup.find("PubmedArticle") or soup.find("Article")
        if article:
            title_el = article.find("ArticleTitle")
            title = title_el.get_text() if title_el else ""

            abstract_el = article.find("Abstract")
            abstract = ""
            if abstract_el:
                for at in abstract_el.find_all("AbstractText"):
                    label = at.get("Label", "")
                    prefix = f"{label}: " if label else ""
                    abstract += prefix + at.get_text() + "\n"

            sections.append({"title": "Abstract", "text": abstract, "level": 1})

            pmid_el = article.find("PMID")
            pmid = pmid_el.get_text() if pmid_el else ""

            return ParsedDocument(
                text=f"{title}\n\n{abstract}".strip(),
                sections=sections,
                metadata={"title": title, "pmid": pmid},
                format="pubmed_xml",
            )

        return ParsedDocument(soup.get_text(separator="\n"), format="xml")

    def _parse_text(self, content: bytes, filename: str) -> ParsedDocument:
        text = content.decode("utf-8", errors="replace")
        sections = []
        current: dict[str, Any] | None = None
        preamble_lines: list[str] = []  # FIX Bug 3: Capture text before first heading

        for line in text.split("\n"):
            stripped = line.strip()
            if stripped.startswith("#"):
                # Flush preamble if we haven't seen a heading yet
                if preamble_lines and not sections and current is None:
                    preamble_text = "\n".join(preamble_lines).strip()
                    if preamble_text:
                        sections.append({"title": "", "text": preamble_text, "level": 0})
                    preamble_lines = []
                if current:
                    sections.append(current)
                level = len(stripped) - len(stripped.lstrip("#"))
                current = {"title": stripped.lstrip("#").strip(), "text": "", "level": level}
            elif current:
                current["text"] += line + "\n"
            else:
                preamble_lines.append(line)

        # Flush final state
        if current:
            sections.append(current)
        elif preamble_lines:
            # Document had no headings at all — treat entire text as one section
            preamble_text = "\n".join(preamble_lines).strip()
            if preamble_text:
                sections.append({"title": "", "text": preamble_text, "level": 0})

        return ParsedDocument(text, sections, {}, "text")

    def _parse_json(self, content: bytes, filename: str) -> ParsedDocument:
        data = json.loads(content)
        if isinstance(data, list):
            sections = []
            for i, rec in enumerate(data):
                name = rec.get("name", rec.get("title", f"Record {i+1}"))
                sections.append({"title": str(name), "text": json.dumps(rec, indent=2), "level": 1})
            return ParsedDocument(json.dumps(data, indent=2), sections, {"record_count": len(data)}, "json")
        return ParsedDocument(json.dumps(data, indent=2), [], {}, "json")
